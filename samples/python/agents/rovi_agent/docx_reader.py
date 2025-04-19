"""
Utility module for reading and parsing DOCX files.
"""
import os
import glob
from typing import Dict, List, Optional
from docx import Document
import re

# Cache for storing parsed content
_concept_cache = {}

def read_docx(file_path: str) -> str:
    """
    Read a DOCX file and return its content as a string.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        The content of the DOCX file as a string.
    """
    try:
        doc = Document(file_path)
        content = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))

        return "\n".join(content)
    except Exception as e:
        return f"Error reading {file_path}: {str(e)}"

def summarize_concept(content: str, max_length: int = 1000) -> str:
    """
    Summarize the concept content to a reasonable length.

    Args:
        content: The full content of the concept.
        max_length: Maximum length of the summary.

    Returns:
        A summarized version of the content.
    """
    # Split content into sections based on newlines
    sections = content.split('\n\n')

    # Extract key information
    summary_parts = []

    # Try to find a title or introduction
    if sections and len(sections) > 0:
        summary_parts.append(sections[0])

    # Look for key sections with important keywords
    key_patterns = [
        r'(?i)mô tả|giới thiệu|tổng quan|overview',
        r'(?i)mục tiêu|objectives|goal',
        r'(?i)hoạt động|activities',
        r'(?i)lợi ích|benefits',
        r'(?i)đối tượng|target',
        r'(?i)thời gian|time|duration',
        r'(?i)địa điểm|location',
        r'(?i)chi phí|cost|pricing'
    ]

    for section in sections[1:]:
        for pattern in key_patterns:
            if re.search(pattern, section):
                summary_parts.append(section)
                break

    # Combine the summary parts
    summary = "\n\n".join(summary_parts)

    # Truncate if too long
    if len(summary) > max_length:
        summary = summary[:max_length] + "..."

    return summary

def get_concept_info(concept_name: Optional[str] = None) -> str:
    """
    Get information about a specific teambuilding concept or list all available concepts.

    Args:
        concept_name: Optional name of the specific concept to retrieve information about.
                     If not provided, returns a list of available concepts.

    Returns:
        Information about the requested teambuilding concept or a list of available concepts.
    """
    try:
        # Path to the teambuilding concept files
        concept_dir = os.path.join(os.path.dirname(__file__), "teambuilding_concept")

        # Get list of available concepts
        concept_files = glob.glob(os.path.join(concept_dir, "*.docx"))

        # Extract concept names from filenames
        concept_names = []
        concept_file_map = {}

        for file_path in concept_files:
            file_name = os.path.basename(file_path)

            # Extract the main concept name (before any dash or period)
            if "Mô tả Concept - " in file_name:
                concept_name_extracted = file_name.replace("Mô tả Concept - ", "").replace(".docx", "")
            else:
                # Handle different file naming patterns
                if " - " in file_name:
                    parts = file_name.split(" - ")
                    concept_name_extracted = parts[0].replace(".docx", "")
                else:
                    concept_name_extracted = file_name.replace(".docx", "")

            concept_names.append(concept_name_extracted)
            concept_file_map[concept_name_extracted] = file_path



        if concept_name is None:
            # Return a formatted list of all available concepts
            concepts_list = "\n".join([f"- {name}" for name in concept_names])
            return f"Các gói teambuilding hiện có:\n{concepts_list}"

        # Find the closest matching concept
        matching_concepts = [c for c in concept_names if concept_name.lower() in c.lower()]

        # If no exact match, try a more flexible matching approach
        if not matching_concepts:
            # Try matching with individual words
            concept_words = concept_name.lower().split()
            for concept in concept_names:
                # Check if any word in the query matches part of the concept name
                if any(word in concept.lower() for word in concept_words):
                    matching_concepts.append(concept)

        if not matching_concepts:
            concepts_list = "\n".join([f"- {name}" for name in concept_names])
            return f"Không tìm thấy gói teambuilding phù hợp với '{concept_name}'. Các gói hiện có:\n{concepts_list}"

        concept = matching_concepts[0]
        concept_file = concept_file_map[concept]

        # Check if we have already parsed this concept
        if concept in _concept_cache:
            return _concept_cache[concept]

        # Read and parse the DOCX file
        content = read_docx(concept_file)

        # Summarize the content
        summary = summarize_concept(content)

        # Format the response
        response = f"Thông tin về gói '{concept}':\n\n{summary}"

        # Cache the response
        _concept_cache[concept] = response

        return response
    except Exception as e:
        import traceback
        print(f"Error in get_concept_info: {str(e)}")
        print(traceback.format_exc())
        return f"Lỗi khi truy xuất thông tin gói teambuilding: {str(e)}"

def get_all_concepts_summary() -> Dict[str, str]:
    """
    Get a summary of all available teambuilding concepts.

    Returns:
        A dictionary mapping concept names to their summaries.
    """
    # Path to the teambuilding concept files
    concept_dir = os.path.join(os.path.dirname(__file__), "teambuilding_concept")

    # Get list of available concepts
    concept_files = glob.glob(os.path.join(concept_dir, "*.docx"))

    # Extract concept names from filenames
    concept_names = []
    concept_file_map = {}

    for file_path in concept_files:
        file_name = os.path.basename(file_path)
        # Extract the main concept name (before any dash or period)
        if "Mô tả Concept - " in file_name:
            concept_name = file_name.replace("Mô tả Concept - ", "").replace(".docx", "")
        else:
            # Handle different file naming patterns
            if " - " in file_name:
                parts = file_name.split(" - ")
                concept_name = parts[0].replace(".docx", "")
            else:
                concept_name = file_name.replace(".docx", "")

        concept_names.append(concept_name)
        concept_file_map[concept_name] = file_path

    # Read and summarize each concept
    summaries = {}
    for concept in concept_names:
        concept_file = concept_file_map[concept]

        # Check if we have already parsed this concept
        if concept in _concept_cache:
            summaries[concept] = _concept_cache[concept]
            continue

        # Read and parse the DOCX file
        content = read_docx(concept_file)

        # Summarize the content (shorter summary for the overview)
        summary = summarize_concept(content, max_length=300)

        # Format the response
        response = f"Thông tin về gói '{concept}':\n{summary}"

        # Cache the response
        _concept_cache[concept] = response

        # Add to summaries
        summaries[concept] = response

    return summaries

def get_concepts_brief_overview() -> str:
    """
    Get a brief overview of all available teambuilding concepts.

    Returns:
        A string containing a brief overview of all available concepts.
    """
    # Path to the teambuilding concept files
    concept_dir = os.path.join(os.path.dirname(__file__), "teambuilding_concept")

    # Get list of available concepts
    concept_files = glob.glob(os.path.join(concept_dir, "*.docx"))

    # Extract concept names from filenames
    concept_names = []
    concept_file_map = {}

    for file_path in concept_files:
        file_name = os.path.basename(file_path)
        # Extract the main concept name (before any dash or period)
        if "Mô tả Concept - " in file_name:
            concept_name = file_name.replace("Mô tả Concept - ", "").replace(".docx", "")
        else:
            # Handle different file naming patterns
            if " - " in file_name:
                parts = file_name.split(" - ")
                concept_name = parts[0].replace(".docx", "")
            else:
                concept_name = file_name.replace(".docx", "")

        concept_names.append(concept_name)
        concept_file_map[concept_name] = file_path

    # Create a brief overview of each concept
    overview_parts = ["Các gói teambuilding hiện có:"]

    for concept in concept_names:
        concept_file = concept_file_map[concept]

        # Read and parse the DOCX file
        content = read_docx(concept_file)

        # Extract the first paragraph or a key sentence
        first_paragraph = content.split('\n\n')[0] if '\n\n' in content else content
        brief = first_paragraph[:150] + "..." if len(first_paragraph) > 150 else first_paragraph

        overview_parts.append(f"- {concept}: {brief}")

    return "\n\n".join(overview_parts)
